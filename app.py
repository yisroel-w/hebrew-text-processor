import streamlit as st
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import base64

def extract_text_from_docx(uploaded_file):
    doc = Document(uploaded_file)
    full_text = []
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)
    return '\n'.join(full_text)

def create_word_doc(text, main_font_size, bold_font_size, direction):
    doc = Document()
    
    # Configure document
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    
    # Process text and extract bold references
    bold_refs = []
    column_text = []
    lines = text.split('\n')
    
    for line in lines:
        while '**' in line:
            start = line.find('**')
            end = line.find('**', start + 2)
            if end != -1:
                bold_refs.append(line[start + 2:end])
                line = line[:start] + line[end + 2:]
        if line.strip():
            column_text.append(line)
    
    # Add bold references
    if bold_refs:
        para = doc.add_paragraph()
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT if direction == "rtl" else WD_ALIGN_PARAGRAPH.LEFT
        run = para.add_run(' '.join(bold_refs))
        run.font.bold = True
        run.font.size = Pt(bold_font_size)
    
    # Create two columns
    section.start_type = 2
    section.column_count = 2
    
    # Split text between columns
    mid_point = len(column_text) // 2
    
    alignment = WD_ALIGN_PARAGRAPH.RIGHT if direction == "rtl" else WD_ALIGN_PARAGRAPH.LEFT
    
    # Add text to columns
    for text in column_text[:mid_point]:
        para = doc.add_paragraph()
        para.paragraph_format.alignment = alignment
        run = para.add_run(text)
        run.font.size = Pt(main_font_size)
    
    doc.add_paragraph().add_run().add_break(docx.enum.text.WD_BREAK.COLUMN)
    
    for text in column_text[mid_point:]:
        para = doc.add_paragraph()
        para.paragraph_format.alignment = alignment
        run = para.add_run(text)
        run.font.size = Pt(main_font_size)
    
    return doc

def get_binary_file_downloader_html(doc):
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()

def main():
    st.set_page_config(page_title="Hebrew Text Layout Processor", layout="wide")
    
    st.title("Hebrew Text Layout Processor")
    
    # Create two columns for layout
    col1, col2 = st.columns([2, 1])
    
    with col1:
        st.subheader("Input Methods")
        
        # File upload
        uploaded_file = st.file_uploader("Upload a Word document", type=['docx'])
        
        if uploaded_file is not None:
            text_content = extract_text_from_docx(uploaded_file)
            st.success("Document uploaded successfully!")
        else:
            text_content = ""
            
        st.subheader("Edit Text")
        text_input = st.text_area(
            "Enter or edit your text (use **text** for bold)",
            value=text_content,
            height=300,
            help="Surround text with ** to make it bold"
        )
        
        st.subheader("Settings")
        col_settings1, col_settings2, col_settings3 = st.columns(3)
        
        with col_settings1:
            main_font_size = st.number_input("Main Text Font Size", 8, 72, 12)
        
        with col_settings2:
            bold_font_size = st.number_input("Bold Text Font Size", 8, 72, 14)
            
        with col_settings3:
            direction = st.selectbox("Text Direction", ["rtl", "ltr"])
    
    with col2:
        st.subheader("Preview")
        if text_input:
            # Show bold references
            bold_refs = []
            preview_text = text_input
            while '**' in preview_text:
                start = preview_text.find('**')
                end = preview_text.find('**', start + 2)
                if end != -1:
                    bold_refs.append(preview_text[start + 2:end])
                    preview_text = preview_text[:start] + preview_text[end + 2:]
            
            st.write("**Bold References:**")
            st.write(", ".join(bold_refs))
            
            st.write("**Column Text:**")
            st.write(preview_text)
    
        st.markdown("---")
        st.subheader("Instructions")
        st.markdown("""
        1. Upload a Word document or enter text directly
        2. Mark text as bold using **double asterisks**
        3. Adjust font sizes and text direction
        4. Click Generate Document to download
        """)
    
    # Generate document
    if st.button("Generate Document"):
        if text_input:
            doc = create_word_doc(text_input, main_font_size, bold_font_size, direction)
            doc_binary = get_binary_file_downloader_html(doc)
            
            st.download_button(
                label="Download Word Document",
                data=doc_binary,
                file_name="hebrew_layout.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        else:
            st.warning("Please enter some text or upload a document first!")

if __name__ == "__main__":
    main()
